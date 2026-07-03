from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


TEMPLATE = Path("/Users/berlin/Library/Containers/com.tencent.xinWeChat/Data/Documents/xwechat_files/wxid_xie0kmelerqp22_0145/msg/file/2026-06/Linux操作系统内核课程报告模板.docx")
OUT = Path("/Users/berlin/Desktop/flagsparse_test-hip/outputs/毕柏林-202522842018586-Linux内核Rust课程报告.docx")
ASSET_DIR = Path("/Users/berlin/Desktop/flagsparse_test-hip/work/course_report/assets")


def set_font(run, name="宋体", size=10.5, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman" if name in ("宋体", "黑体", "楷体_GB2312") else name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman" if name in ("宋体", "黑体", "楷体_GB2312") else name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def para(doc, text="", align=None, first_line=True, before=0, after=0, line=1.5, font="宋体", size=10.5):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if first_line:
        fmt.first_line_indent = Pt(21)
    r = p.add_run(text)
    set_font(r, font, size)
    return p


def mixed_para(doc, parts, align=None, first_line=True, before=0, after=0, line=1.5):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if first_line:
        fmt.first_line_indent = Pt(21)
    for text, font, size, bold in parts:
        r = p.add_run(text)
        set_font(r, font, size, bold=bold)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(9 if level == 1 else 6)
    fmt.space_after = Pt(6)
    fmt.line_spacing = 1.5
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        set_font(r, "黑体", 15, bold=False)
    else:
        r = p.add_run(text)
        set_font(r, "黑体", 12, bold=False)
    return p


def cover_line(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.left_indent = Cm(3.2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p.add_run(label)
    set_font(r1, "宋体", 14)
    r2 = p.add_run(value)
    set_font(r2, "宋体", 14)
    return p


def clear_body(doc):
    body = doc._body._element
    sect_pr = body.sectPr
    if sect_pr is not None:
        sect_pr = deepcopy(sect_pr)
    for child in list(body):
        body.remove(child)
    if sect_pr is not None:
        body.append(sect_pr)


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    set_font(run, "宋体", 9)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)


def font_path():
    candidates = [
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return path
    return None


def draw_center(draw, box, text, font, fill=(31, 55, 95)):
    left, top, right, bottom = box
    lines = text.split("\n")
    line_heights = []
    widths = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        widths.append(bbox[2] - bbox[0])
        line_heights.append(bbox[3] - bbox[1])
    total_h = sum(line_heights) + (len(lines) - 1) * 8
    y = top + (bottom - top - total_h) / 2
    for line, w, h in zip(lines, widths, line_heights):
        x = left + (right - left - w) / 2
        draw.text((x, y), line, font=font, fill=fill)
        y += h + 8


def arrow(draw, start, end, fill=(70, 86, 115), width=4):
    draw.line([start, end], fill=fill, width=width)
    sx, sy = start
    ex, ey = end
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        points = [(ex, ey), (ex - 16 * direction, ey - 9), (ex - 16 * direction, ey + 9)]
    else:
        direction = 1 if ey > sy else -1
        points = [(ex, ey), (ex - 9, ey - 16 * direction), (ex + 9, ey - 16 * direction)]
    draw.polygon(points, fill=fill)


def rounded_box(draw, box, text, font, fill, outline=(94, 111, 140), text_fill=(20, 33, 61)):
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=3)
    draw_center(draw, box, text, font, text_fill)


def create_figures():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    fp = font_path()
    title_font = ImageFont.truetype(fp, 34) if fp else ImageFont.load_default()
    mid_font = ImageFont.truetype(fp, 25) if fp else ImageFont.load_default()
    small_font = ImageFont.truetype(fp, 20) if fp else ImageFont.load_default()

    # Figure 1: safety design space.
    img = Image.new("RGB", (1400, 720), "white")
    d = ImageDraw.Draw(img)
    d.text((50, 35), "内核安全路径对比", font=title_font, fill=(11, 37, 69))
    boxes = [
        ((70, 160, 380, 340), "硬件隔离\n进程 / 地址空间\n安全强但切换开销高", (232, 238, 247)),
        ((545, 160, 855, 340), "GC 类型安全语言\n运行时复杂\n时延不可预测", (245, 241, 228)),
        ((1020, 160, 1330, 340), "Rust 内核\n无 GC + 类型安全\nunsafe 边界可审计", (230, 244, 237)),
    ]
    for box, text, fill in boxes:
        rounded_box(d, box, text, mid_font, fill)
    arrow(d, (380, 250), (545, 250))
    arrow(d, (855, 250), (1020, 250))
    d.text((120, 450), "论文主张：不是在内核中完全放弃底层控制，而是把风险收敛到少量可信抽象。", font=mid_font, fill=(37, 52, 87))
    d.rounded_rectangle((120, 515, 1280, 630), radius=16, fill=(248, 250, 252), outline=(210, 218, 230), width=2)
    d.text((165, 550), "可信代码基：libcore 中必要 unsafe + 内核底层 unsafe 封装；上层逻辑尽量保持 safe Rust。", font=small_font, fill=(75, 85, 99))
    fig1 = ASSET_DIR / "kernel_safety_paths.png"
    img.save(fig1)

    # Figure 2: event-driven SimpleRng sharing.
    img = Image.new("RGB", (1400, 820), "white")
    d = ImageDraw.Draw(img)
    d.text((50, 35), "事件驱动内核中的共享状态冲突", font=title_font, fill=(11, 37, 69))
    rounded_box(d, (80, 180, 360, 330), "用户进程\nProcess", mid_font, (232, 238, 247))
    rounded_box(d, (520, 180, 850, 330), "系统调用路径\ncommand() 设置 busy", mid_font, (239, 246, 255))
    rounded_box(d, (1040, 180, 1320, 330), "SimpleRng\n共享可变状态", mid_font, (230, 244, 237))
    rounded_box(d, (520, 520, 850, 670), "硬件中断回调\ndeliver() 清除 busy", mid_font, (255, 247, 237))
    rounded_box(d, (80, 520, 360, 670), "RNG 硬件\n随机数完成", mid_font, (245, 241, 228))
    arrow(d, (360, 255), (520, 255))
    arrow(d, (850, 255), (1040, 255))
    arrow(d, (360, 595), (520, 595))
    arrow(d, (850, 595), (1040, 330))
    d.text((995, 455), "Rust 需要证明：\n同一时刻只有一个可变访问者", font=small_font, fill=(121, 65, 20))
    d.text((115, 735), "Cell 适合 busy 这类小状态；TakeCell 适合复杂结构的“取出-使用-归还”。", font=mid_font, fill=(37, 52, 87))
    fig2 = ASSET_DIR / "simple_rng_paths.png"
    img.save(fig2)

    # Figure 3: Linux Rust module boundary.
    img = Image.new("RGB", (1400, 820), "white")
    d = ImageDraw.Draw(img)
    d.text((50, 35), "Rust for Linux 的边界思路", font=title_font, fill=(11, 37, 69))
    rounded_box(d, (90, 150, 390, 300), "安全 Rust 驱动\n模块业务逻辑", mid_font, (230, 244, 237))
    rounded_box(d, (550, 150, 850, 300), "Rust 内核抽象\n锁 / 引用 / 设备对象", mid_font, (232, 238, 247))
    rounded_box(d, (1010, 150, 1310, 300), "unsafe 封装层\nFFI / MMIO / DMA", mid_font, (255, 247, 237))
    rounded_box(d, (550, 500, 850, 650), "Linux C 内核\n调度 / 内存 / VFS", mid_font, (248, 250, 252))
    rounded_box(d, (1010, 500, 1310, 650), "硬件设备\n寄存器 / 中断", mid_font, (245, 241, 228))
    arrow(d, (390, 225), (550, 225))
    arrow(d, (850, 225), (1010, 225))
    arrow(d, (1160, 300), (1160, 500))
    arrow(d, (1010, 575), (850, 575))
    d.text((125, 405), "目标：让多数新驱动代码停留在 safe Rust，\n把不可避免的底层风险集中到小而稳定的接口。", font=mid_font, fill=(37, 52, 87))
    fig3 = ASSET_DIR / "rust_for_linux_boundary.png"
    img.save(fig3)
    return fig1, fig2, fig3


def add_caption(doc, text):
    p = para(doc, text, WD_ALIGN_PARAGRAPH.CENTER, False, before=2, after=6, line=1.2, font="宋体", size=9)
    for run in p.runs:
        run.italic = True
    return p


def add_figure(doc, path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(14.6))
    add_caption(doc, caption)


def code_block(doc, code, caption=None):
    if caption:
        mixed_para(doc, [("代码：", "黑体", 10.5, True), (caption, "宋体", 10.5, False)], first_line=False, before=4, after=2, line=1.2)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.right_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.05
    for idx, line_text in enumerate(code.strip("\n").splitlines()):
        if idx:
            p.add_run().add_break()
        r = p.add_run(line_text)
        set_font(r, "Menlo", 8.5)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F4F6F8")
    p._p.get_or_add_pPr().append(shading)
    border = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "8")
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), "8AA4C8")
    border.append(left)
    p._p.get_or_add_pPr().append(border)


def comparison_table(doc):
    table = doc.add_table(rows=1, cols=4)
    table.autofit = False
    hdr = table.rows[0].cells
    headers = ["维度", "C 语言内核常见风险", "Rust 对应机制", "在 Linux 内核中的意义"]
    widths = [2.2, 4.2, 3.6, 4.6]
    for i, cell in enumerate(hdr):
        cell.text = headers[i]
        cell.width = Cm(widths[i])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                set_font(r, "黑体", 9.5, bold=True)
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "E8EEF5")
        tc_pr.append(shd)
    rows = [
        ("生命周期", "手动释放或引用计数错误导致 use-after-free。", "所有权、生命周期、Drop 自动释放。", "降低驱动对象、文件对象、缓冲区等资源泄漏或过早释放概率。"),
        ("别名与可变性", "多个指针同时写同一对象，边界依赖人工约定。", "可变借用唯一性，Cell/TakeCell 封装内部可变性。", "把共享状态的访问路径显式化，便于审计中断和回调。"),
        ("错误处理", "返回码容易被忽略，空指针检查不完整。", "Result/Option 强制表达失败和空值。", "让初始化失败、设备不存在、缓冲区非法等路径更难被遗漏。"),
        ("硬件访问", "整数地址和位操作自由度高，误写寄存器风险大。", "类型化寄存器、枚举和安全 wrapper。", "适合 MMIO、DMA、USB/PCI 等硬件协议边界。"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            cells[i].width = Cm(widths[i])
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.15
                if i == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    set_font(r, "宋体", 9)
    add_caption(doc, "表 1  C 与 Rust 在内核开发中的风险表达方式对比")


def build():
    fig_safety, fig_rng, fig_linux = create_figures()
    doc = Document(str(TEMPLATE))
    clear_body(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)
    add_page_number(section)

    # Cover page, following the supplied course-report template.
    for _ in range(5):
        para(doc, "", first_line=False, line=1.0)
    p = para(doc, "《Linux操作系统内核》", WD_ALIGN_PARAGRAPH.CENTER, False, line=1.0, font="宋体", size=18)
    p.runs[0].bold = False
    para(doc, "2026年春季课程报告", WD_ALIGN_PARAGRAPH.CENTER, False, line=1.0, font="宋体", size=18)
    for _ in range(2):
        para(doc, "", first_line=False, line=1.0)
    para(doc, "用 Rust 编写操作系统内核的理由", WD_ALIGN_PARAGRAPH.CENTER, False, line=1.0, font="宋体", size=16)
    for _ in range(4):
        para(doc, "", first_line=False, line=1.0)
    cover_line(doc, "论文题目：", "The Case for Writing a Kernel in Rust")
    cover_line(doc, "学生姓名：", "毕柏林")
    cover_line(doc, "学    号：", "202522842018586")
    cover_line(doc, "专业班级：", "")
    cover_line(doc, "完成时间：", "2026年7月1日")

    p = para(doc, "摘要", WD_ALIGN_PARAGRAPH.CENTER, False, before=3, line=1.5, font="黑体", size=15)
    p.paragraph_format.page_break_before = True
    p.paragraph_format.space_after = Pt(6)
    para(
        doc,
        "操作系统内核长期依赖 C 语言与硬件地址空间隔离来获得性能和控制力，但这种传统路线也带来了缓冲区溢出、悬垂指针、越界访问和驱动错误扩大化等安全问题。"
        "论文 The Case for Writing a Kernel in Rust 从嵌入式内核实践出发，论证了在不依赖垃圾回收运行时的前提下，Rust 可以通过所有权、借用检查和类型系统为内核提供接近 C 语言的执行效率与更强的内存安全保障。"
        "本文在论文和汇报材料的基础上，结合 Linux 内核的模块化结构、驱动模型、内存管理和 Rust for Linux 的工程方向，分析 Rust 内核开发的技术价值、可信代码基边界以及仍需面对的多核并发、硬件抽象和生态迁移问题。"
        "报告认为，Rust 的意义并不是完全消除 unsafe，而是将 unsafe 压缩到少量可审计的底层抽象中，使上层内核逻辑能够在更严格的类型约束下表达。",
    )
    mixed_para(doc, [("关键词：", "黑体", 10.5, True), ("Linux内核；Rust；内存安全；可信代码基；TakeCell；驱动程序", "宋体", 10.5, False)], first_line=False)
    para(doc, "", first_line=False)
    para(doc, "Abstract", WD_ALIGN_PARAGRAPH.CENTER, False, line=1.5, font="Times New Roman", size=15)
    para(
        doc,
        "Traditional operating-system kernels rely heavily on C and hardware-enforced isolation to obtain performance and low-level control, but this design also exposes kernels to buffer overflows, dangling pointers, invalid memory access, and driver-induced failures. "
        "Based on the paper The Case for Writing a Kernel in Rust and the accompanying presentation, this report explains why Rust is a promising language for kernel development without requiring a garbage-collected runtime. "
        "It further connects the paper's core ideas with Linux kernel mechanisms, including driver development, memory management, interrupt handling, and the recent Rust-for-Linux direction. "
        "The central argument is that Rust does not remove all unsafe operations from a kernel; instead, it makes the boundary of unsafe code smaller, clearer, and easier to audit, while allowing most kernel logic to remain memory safe.",
        first_line=True,
        font="Times New Roman",
    )
    mixed_para(doc, [("Key words: ", "Times New Roman", 10.5, True), ("Linux kernel; Rust; memory safety; trusted computing base; TakeCell; device driver", "Times New Roman", 10.5, False)], first_line=False)

    sections = [
        (
            "1. 绪论",
            [
                "Linux 内核是现代操作系统中最核心的软件层，承担进程调度、内存管理、文件系统、网络协议栈、设备驱动和安全访问控制等职责。由于内核运行在高特权级，一旦内核代码发生越界写、释放后使用或空指针解引用，影响往往不局限于单个应用程序，而可能导致系统崩溃、权限提升或长期潜伏的安全漏洞。",
                "传统内核长期选择 C 语言，是因为 C 提供了接近硬件的表达能力、可预测的性能和成熟的生态。然而，C 语言把内存生命周期、指针别名、并发同步等大量责任交给程序员。对于由数千万行代码构成、包含大量第三方驱动的 Linux 内核而言，这种模式使安全性高度依赖人工审查、测试覆盖和运行时防护。",
                "论文《The Case for Writing a Kernel in Rust》提出了一个值得操作系统课程关注的问题：能否使用一种既具备系统级性能，又能在编译期提供内存安全保证的语言来编写内核？作者选择 Rust 的原因在于，Rust 没有垃圾回收器，能够保留底层内存布局控制，同时通过所有权和借用检查避免大量传统内存错误。"
            ],
        ),
        (
            "2. 传统内核安全机制的局限",
            [
                "操作系统通常通过进程抽象和虚拟地址空间隔离用户程序。进程边界由硬件 MMU、页表和特权级机制共同维护，因此一个用户进程的错误通常不会直接破坏另一个进程或内核数据结构。但在内核内部，许多代码共享同一地址空间，传统设计默认内核代码彼此可信。驱动程序、文件系统和网络协议栈中的一次非法写入，都可能越过模块边界影响整个内核。",
                "已有研究尝试使用微内核、轻量级上下文或隔离驱动框架降低内核内部故障传播范围。这些方案能够提高可靠性，却常常需要额外的地址空间切换、消息传递或上下文保存。对于高频调用的内核路径，过高的隔离开销会影响系统吞吐与延迟，使其难以全面替代单体内核中的直接函数调用。",
                "另一条路线是使用类型安全语言构建内核或可扩展内核。论文提到 Spin 和 Singularity 等系统证明了语言级安全的潜力，但它们依赖带垃圾回收的运行时。垃圾回收会引入不可预测的暂停、后台锁竞争和内存布局限制，而内核恰恰需要对时序、内存位置和资源释放拥有精确控制。因此，适合内核的语言需要同时满足类型安全、低运行时负担和硬件可控性。"
            ],
        ),
        (
            "3. Rust 内存模型及其内核编程冲突",
            [
                "Rust 的核心机制是所有权、生命周期和借用检查。每个值都有唯一所有者，当所有者离开作用域时资源自动释放；同一时刻要么存在一个可变引用，要么存在多个不可变引用。这套规则在编译期排除了释放后使用、重复释放和数据竞争等常见问题，也使 Rust 在没有垃圾回收器的情况下获得自动内存管理能力。",
                "但是，内核程序天然具有事件驱动和共享状态的特点。以论文中的随机数生成器为例，系统调用路径需要调用 SimpleRng.command() 将 busy 状态置为 true，硬件中断回调路径又需要调用 SimpleRng.deliver() 将 busy 状态置为 false。两个异步上下文都需要修改同一个对象，直接触碰了 Rust “同一时间最多一个可变引用”的限制。",
                "这种冲突并不说明 Rust 不适合内核，而是说明内核需要把共享可变状态重新建模。C 语言可以通过裸指针或全局变量轻易表达多个可变别名，但这种自由也正是类型混淆、悬垂指针和越界写入的来源。Rust 的策略是先拒绝无法证明安全的表达，再要求开发者通过受约束的抽象重新给出安全接口。"
            ],
        ),
        (
            "4. 面向 Rust 内核的最小可信抽象",
            [
                "论文将可信代码分为两类。第一类是 Rust 语言和核心库中必要的 unsafe 实现，例如数组边界检查、迭代器优化、编译器内置操作、基础类型转换和 Cell 等内部可变性机制。内核通常只需要 libcore，若需要通用动态分配才进一步引入 liballoc，这比依赖完整标准库或垃圾回收运行时更轻量。",
                "第二类是内核自身必须实现的底层 unsafe 代码，包括上下文切换、系统调用陷入、中断/异常处理、内存映射 I/O、用户空间缓冲区检查、内存分配器以及 TakeCell 等抽象。它们之所以可信，是因为底层机制无法完全由安全 Rust 表达；它们之所以可控，是因为 unsafe 被封装在少量模块中，上层调用者只能接触安全接口。",
                "Cell 适合 bool、整数等 Copy 小对象，可通过不可变引用完成内部状态更新，解决 SimpleRng.busy 这类轻量共享状态。TakeCell 则面向复杂结构，它通过 map 闭包暂时取出内部值，闭包结束后再归还，从而保证任意时刻只有一个可变访问者。与互斥锁相比，TakeCell 在资源已被取出时选择不执行闭包而不是阻塞，适合单核嵌入式内核中的非阻塞事件路径。"
            ],
        ),
        (
            "5. 与 Linux 内核机制的结合分析",
            [
                "Linux 内核的设备驱动是内存安全风险最集中的区域之一。驱动需要操作 MMIO 寄存器、DMA 缓冲区、中断回调和硬件描述符，既靠近硬件，又经常由不同厂商维护。Rust 的类型系统可以把“寄存器只能写合法位”“DMA 缓冲区必须在操作完成前保持有效”“数组长度必须匹配硬件端点数量”等规则编码进类型和生命周期，减少约定式编程带来的错误。",
                "在 Linux 内核内存管理中，页表、slab 分配器、引用计数和用户态缓冲区拷贝都涉及明确的所有权边界。Rust 无法替代硬件页表和内核分配器本身，但可以帮助模块作者在接口层表达资源归属，例如用 RAII 方式保证锁释放，用生命周期限制缓冲区有效期，用 Result 和 Option 显式处理错误与空值，从而降低遗漏检查的概率。",
                "Rust for Linux 的工程方向与论文观点是一致的：不是把既有 Linux 内核一次性重写成 Rust，而是在可控边界内逐步允许 Rust 模块进入内核，先从驱动和较新子系统积累经验。对于大型通用内核，现实路径更可能是 C 与 Rust 长期共存，Rust 负责新模块和高风险接口封装，C 代码继续承担成熟核心路径。"
            ],
        ),
        (
            "6. 论文案例研究",
            [
                "在 DMA 场景中，硬件能够绕过 CPU 直接访问内存，因此错误地址或错误长度会造成严重破坏。论文用 &'static mut [u8] 表示 DMA 缓冲区，slice 携带长度信息以防越界，'static 生命周期保证 DMA 完成前缓冲区不会提前释放，TakeCell 保证缓冲区可变访问的唯一性。这样，原本依赖程序员约定的安全条件被转化为编译器可检查的类型约束。",
                "在 USB 场景中，硬件描述符和控制寄存器要求严格的内存布局。Rust 可以用枚举限制控制位，用固定长度数组表达端点数量，用结构体映射硬件寄存器。相比 C 语言中随意的整数位操作，这种方式把协议规则沉淀为类型规则，在不增加运行时开销的前提下减少误写寄存器和数组越界的风险。",
                "对于 buffer cache、页表、链表和树等复杂结构，Rust 面临循环引用和多别名问题。论文给出的思路是尽量使用逻辑标识代替真实指针，例如设备号和扇区号；必须存在双向引用时，再用 Cell 与 TakeCell 等抽象封装共享可变状态。这为 Linux 内核中复杂数据结构的 Rust 化提供了启示：关键不是机械翻译指针，而是重新设计所有权边界。"
            ],
        ),
        (
            "7. 局限性与个人评价",
            [
                "论文的实证对象主要是低功耗、单核、嵌入式场景，并没有完整评估大型多核通用操作系统。Linux 内核还需要面对抢占、RCU、细粒度锁、NUMA、缓存一致性和海量驱动生态等复杂问题。TakeCell 的非阻塞语义适合某些单核事件路径，但在多核环境中仍需要与锁、原子操作、引用计数和内存屏障等机制结合。",
                "Rust 也不能自动保证内核逻辑正确。死锁、权限检查遗漏、协议状态机错误、侧信道问题和错误的硬件规格理解，仍然可能出现在安全 Rust 代码中。因此，Rust 更准确的价值是缩小内存安全类漏洞的空间，而不是替代测试、形式化验证、代码审查和运行时防护。",
                "从课程学习角度看，这篇论文的意义在于提供了理解现代内核安全的新视角。传统操作系统课程强调硬件隔离、系统调用、页表和进程模型，而 Rust 进一步把安全边界推进到语言和类型系统层面。它提示我们，内核安全不仅是运行时机制问题，也是编程语言、抽象设计和工程组织共同作用的结果。"
            ],
        ),
        (
            "8. 结论",
            [
                "《The Case for Writing a Kernel in Rust》并未主张内核可以完全没有 unsafe，而是说明 unsafe 应该被限制在最小、清晰、可审计的可信代码基中。Rust 的所有权、借用检查、生命周期和无 GC 设计，使它有机会在保持底层控制力的同时提升内核内存安全。",
                "结合 Linux 内核来看，Rust 最适合从驱动程序、硬件抽象层和新子系统接口逐步切入。它能够把许多隐含约定转化为编译期约束，降低驱动错误破坏内核全局状态的概率。未来 Rust 在 Linux 中的价值，将取决于语言抽象、内核基础设施、开发者生态和长期维护成本之间的平衡。"
            ],
        ),
    ]

    for title, paragraphs in sections:
        h = heading(doc, title, 1)
        if title.startswith("1."):
            h.paragraph_format.page_break_before = True
        for text in paragraphs:
            para(doc, text)
        if title.startswith("1."):
            add_figure(doc, fig_safety, "图 1  从硬件隔离、GC 类型安全语言到 Rust 内核的设计路线")
            para(
                doc,
                "本报告不仅复述论文提出的 Rust 内核可行性论证，还把它放回 Linux 内核课程的知识体系中理解。"
                "具体来说，报告从“内核为什么危险”出发，先分析硬件隔离和 C 语言开发模式的边界，再讨论 Rust 如何通过编译期规则限制指针别名和资源生命周期，"
                "最后联系 Linux 驱动、DMA、MMIO、用户缓冲区和模块加载等机制，说明 Rust 适合作为增量式安全改造工具，而不是简单替代整个内核。"
            )
        elif title.startswith("2."):
            comparison_table(doc)
        elif title.startswith("3."):
            add_figure(doc, fig_rng, "图 2  SimpleRng 在系统调用路径和硬件回调路径之间形成共享状态")
            para(
                doc,
                "对于 busy 这样的布尔状态，如果仍要求所有调用路径都持有 &mut self，内核事件模型会很难表达。"
                "Cell 的作用是把“外部不可变引用”和“内部值更新”隔离开：调用者无法取得内部字段的可变引用，只能通过 set/get 复制小值。"
                "这使得多个组件可以共享对 SimpleRng 的不可变引用，而不会形成裸指针式的任意别名。"
            )
            code_block(
                doc,
                """
use core::cell::Cell;

struct SimpleRng {
    busy: Cell<bool>,
}

impl SimpleRng {
    fn command(&self) {
        self.busy.set(true);
        // 启动硬件随机数生成操作
    }

    fn deliver(&self, value: u32) {
        self.busy.set(false);
        // 将随机数返回给等待的系统调用
        let _ = value;
    }
}
""",
                "用 Cell 改写 SimpleRng 的 busy 状态",
            )
        elif title.startswith("4."):
            para(
                doc,
                "TakeCell 解决的是 Cell 无法高效处理大型结构的问题。Cell 需要复制值，适合 bool 或整数；"
                "而内核中的进程控制块、驱动状态、DMA 缓冲区描述符等结构往往较大，复制成本高，也可能不具备 Copy 语义。"
                "TakeCell 通过闭包把可变访问限制在一个短小作用域中，编译器和容器共同保证访问结束后资源被放回。"
            )
            code_block(
                doc,
                """
struct AppState {
    count: u32,
    tx_ready: bool,
}

struct Driver {
    app: TakeCell<AppState>,
}

impl Driver {
    fn on_interrupt(&self) {
        self.app.map(|state| {
            state.count += 1;
            state.tx_ready = true;
        });
    }
}
""",
                "TakeCell 的“取出-使用-归还”模式示意",
            )
            para(
                doc,
                "这段代码的重点不在于语法本身，而在于访问协议：闭包运行期间，AppState 不再同时暴露给其他路径；"
                "闭包结束后，状态重新回到容器中。论文展示的 ARM 汇编说明，这一模式最终可编译为少量加载、置空、判断和回写指令，不需要复杂运行时。"
            )
        elif title.startswith("5."):
            add_figure(doc, fig_linux, "图 3  Rust for Linux 中 safe Rust、unsafe 封装层与 C 内核之间的边界")
            para(
                doc,
                "在 Linux 中引入 Rust 的现实方式，是让 Rust 模块通过内核提供的抽象与现有 C 子系统交互。"
                "下面的代码是 Rust for Linux 风格的模块骨架示意：初始化失败用 Result 返回，模块卸载时通过 Drop 回收资源，日志输出由内核宏完成。"
                "这种写法体现了 Rust 的 RAII 思想，也让资源申请和释放更容易被局部审查。"
            )
            code_block(
                doc,
                """
use kernel::prelude::*;

module! {
    type: RustDemo,
    name: "rust_demo",
    author: "Bibolin",
    description: "A minimal Rust kernel module example",
    license: "GPL",
}

struct RustDemo;

impl kernel::Module for RustDemo {
    fn init(_module: &'static ThisModule) -> Result<Self> {
        pr_info!("Rust kernel module loaded\\n");
        Ok(RustDemo)
    }
}

impl Drop for RustDemo {
    fn drop(&mut self) {
        pr_info!("Rust kernel module unloaded\\n");
    }
}
""",
                "Rust for Linux 模块骨架示意",
            )
            para(
                doc,
                "需要强调的是，Rust 模块仍然运行在内核态，仍然可能因为逻辑错误影响系统；"
                "但它把空指针、生命周期和资源释放这些问题从“人工约定”转向“类型和接口约束”，这正是论文思想在 Linux 生态中的延伸。"
            )
        elif title.startswith("6."):
            para(
                doc,
                "DMA 的核心难点是 CPU 与设备同时观察同一块内存。C 语言中常见写法是把物理地址、长度和控制位分别传递给设备，"
                "但这些值之间的关系并不总能被编译器检查。Rust 更倾向于把缓冲区指针与长度绑定为 slice，并用生命周期说明这块内存在 DMA 完成前必须保持有效。"
            )
            code_block(
                doc,
                """
struct DmaChannel {
    enabled: Cell<bool>,
    buffer: TakeCell<&'static mut [u8]>,
}

impl DmaChannel {
    fn start(&self) {
        self.buffer.map(|buf| {
            let addr = buf.as_mut_ptr();
            let len = buf.len();
            // unsafe 封装层把 addr/len 写入 DMA 寄存器
            let _ = (addr, len);
            self.enabled.set(true);
        });
    }
}
""",
                "DMA 缓冲区的生命周期和长度由类型携带",
            )
            para(
                doc,
                "这类代码仍然需要一个底层 unsafe wrapper 写硬件寄存器，但 unsafe 的职责被缩小为“把已经验证过的地址和长度交给设备”。"
                "上层驱动无需直接拼接整数地址，也不容易把临时栈缓冲区误交给长时间运行的 DMA 操作。"
            )

    heading(doc, "参考文献", 1)
    refs = [
        "Amit Levy, Bradford Campbell, Branden Ghena, Pat Pannuto, Prabal Dutta, Philip Levis. The Case for Writing a Kernel in Rust. Proceedings of APSys '17, 2017. DOI: 10.1145/3124680.3124717.",
        "Amit Levy, Michael P. Andersen, Bradford Campbell, David Culler, Prabal Dutta, Branden Ghena, Philip Levis, Pat Pannuto. Ownership is Theft: Experiences Building an Embedded OS in Rust. Proceedings of PLOS '15, 2015, pp. 21-26.",
        "N. D. Matsakis, F. S. Klock II. The Rust Language. Proceedings of HILT '14, 2014, pp. 103-104.",
        "The kernel development community. Rust - The Linux Kernel documentation. https://docs.kernel.org/rust/index.html.",
        "Rust for Linux Project. Rust for Linux. https://rust-for-linux.com/.",
    ]
    for idx, text in enumerate(refs, 1):
        para(doc, f"[{idx}] {text}", first_line=False, line=1.3)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    build()
